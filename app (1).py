#This system was created by Deliena Tasha Binti Abdul Rahim xdeliena on GitHub

import os, sys, shutil, zipfile, uuid, re
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import gradio as gr
import pandas as pd
from docx import Document
import os, tempfile, zipfile
from docx.shared import Inches 
from huggingface_hub import HfApi
import getpass, requests
from supabase import create_client, Client
from transformers import pipeline

# -------------------------
# Config
# -------------------------
SUPABASE_URL = os.getenv("SUPABASE_URL","XXXXXX")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "XXX") or os.getenv("SUPABASE_KEY","XXX")
if not SUPABASE_URL or not SUPABASE_KEY:
    raise RuntimeError("SUPABASE_URL and a SUPABASE key must be set in env vars")
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
os.makedirs(TEMPLATES_DIR, exist_ok=True)
HF_SPACE_REPO = os.getenv("SPACE_ID") or os.getenv("HF_SPACE_REPO") or "unknown/space"
CACHED_DATA: List[Dict[str, str]] = []
CACHED_COLUMNS: List[str] = []
print(f"🚀 Running in Space: {HF_SPACE_REPO}")

# -------------------------
# Helpers
# -------------------------
def list_templates() -> List[str]:
    try:
        response = supabase.table("templates").select("filename").execute()
        return sorted([r["filename"] for r in response.data])
    except Exception as e:
        print("⚠️ Error listing templates:", e)
        return []

def extract_placeholders(template_name: str) -> List[str]:
    if not template_name:
        return []
    path = get_template_path_from_supabase(template_name)
    if not path or not os.path.exists(path):
        print(f"⚠️ Could not fetch template {template_name} from database")
        return []
    doc = Document(path)
    text = " ".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += " " + " ".join(p.text for p in cell.paragraphs)
    matches = re.findall(r"\{\{(.*?)\}\}|\{(.*?)\}", text)
    return list({(m[0] or m[1]).strip() for m in matches if (m[0] or m[1]).strip()})

def sanitize_filename(name: str) -> str:
    return re.sub(r"[\\/*?<>|:\"\n\r\t]+", "_", name.strip())[:200]

def replace_placeholders(doc: Document, fields: Dict[str, str]) -> Document:
    def process_paragraph(par):
        for run in par.runs:
            text = run.text
            for k, v in fields.items():
                # --- Handle image placeholders ---
                if k.endswith("image") and v and os.path.exists(v):
                    if f"{{{k}}}" in text or f"{{{{{k}}}}}" in text:
                        run.text = text.replace(f"{{{k}}}", "").replace(f"{{{{{k}}}}}", "")
                        new_run = par.add_run()
                        new_run.add_picture(v, width=Inches(1.5))
                        text = run.text
                
                # --- Handle text placeholders ---
                else:
                    if f"{{{k}}}" in text or f"{{{{{k}}}}}" in text or f"{{{k.upper()}}}" in text:
                        val = str(v)
                        if f"{{{k.upper()}}}" in text:
                            val = val.upper()
                        run.text = (
                            text.replace(f"{{{k}}}", val)
                                .replace(f"{{{{{k}}}}}", val)
                                .replace(f"{{{k.upper()}}}", val)
                        )
                        
    for p in doc.paragraphs:
        process_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
    return doc

def generate_single_docx(template_name: str, fields: Dict[str, str], rename_pattern: Optional[str]) -> str:
    tpl_path = get_template_path_from_supabase(template_name)
    if not tpl_path or not os.path.exists(tpl_path):
        raise FileNotFoundError(f"Template {template_name} not found in database")
    doc = Document(tpl_path)
    
    # Normalize field keys to lowercase before replacement
    lower_fields = {k.lower(): v for k, v in fields.items()}
    doc = replace_placeholders(doc, lower_fields)
    base = os.path.splitext(template_name)[0]

    if rename_pattern:
        name = rename_pattern
        # Make matching case-insensitive
        for k, v in fields.items():
            pattern = re.compile(rf"\{{{{\s*{re.escape(k)}\s*\}}}}|\{{\s*{re.escape(k)}\s*\}}", re.IGNORECASE)
            name = pattern.sub(str(v), name)
        name = sanitize_filename(name) or f"{base}_{uuid.uuid4().hex[:6]}"
    else:
        name = f"{base}_{fields.get('name', uuid.uuid4().hex[:6])}"

    # Save temporarily (so user downloads instead of system saving)
    tmp_dir = tempfile.mkdtemp()
    out_path = os.path.join(tmp_dir, f"{name}.docx")
    doc.save(out_path)
    return out_path

# -------------------------
# Data parsing
# -------------------------
def parse_file(path: str) -> List[Dict[str, str]]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        df = pd.read_csv(path, keep_default_na=False)
    else:
        df = pd.read_excel(path, keep_default_na=False)

    # Normalize column names: lowercase + replace spaces with underscores
    df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]

    # Convert each cell, with special handling for dates
    def format_cell(x):
        if pd.isna(x):
            return ""
        if isinstance(x, (datetime, pd.Timestamp)):
            return x.strftime("%#d %B %Y") if os.name == "nt" else x.strftime("%-d %B %Y")
        return str(x).strip()

    records = df.applymap(format_cell).to_dict(orient="records")
    return records

def parse_pasted_text(text: str) -> Tuple[List[Dict[str, str]], List[str]]:
    rows, errors = [], []
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for i, line in enumerate(lines, 1):
        parts = [p.strip() for p in line.split(",") if p.strip()]
        row = {}
        for part in parts:
            if ":" not in part:
                errors.append(f"Line {i}: '{part}' missing ':'")
                continue
            k, v = part.split(":", 1)
            row[k.strip()] = v.strip()  # <-- no normalization, keep exact format
        if row:
            rows.append(row)
    return rows, errors

def get_template_path_from_supabase(filename: str) -> Optional[str]:
    """Download a template file temporarily from Supabase Storage."""
    try:
        if not filename:
            return None
        # Get file URL from table (database)
        res = supabase.table("templates").select("file_url").eq("filename", filename).execute()
        if not res.data:
            print(f"⚠️ Template not found in database table: {filename}")
            return None
        file_url = res.data[0]["file_url"]
        # Download the .docx file into a temporary path
        tmp_path = os.path.join(tempfile.gettempdir(), filename)
        r = requests.get(file_url)
        if r.status_code == 200:
            with open(tmp_path, "wb") as f:
                f.write(r.content)
            return tmp_path
        print(f"⚠️ Failed to download from database: HTTP {r.status_code}")
        return None
    except Exception as e:
        print("❌ Error fetching template:", e)
        return None

# -------------------------
# Template Handlers
# -------------------------
def handle_upload(upload):
    if not upload:
        return gr.update(), gr.update(), gr.update(), gr.update(), "⚠️ No file uploaded"

    try:
        temp_path = upload.name if hasattr(upload, "name") else str(upload)
        filename = os.path.basename(temp_path)

        if not filename.lower().endswith(".docx"):
            return gr.update(), gr.update(), gr.update(), gr.update(), "❌ Only .docx allowed"

        # Upload to Supabase Storage
        with open(temp_path, "rb") as f:
            data = f.read()
        res = supabase.storage.from_("templates").upload(f"templates/{filename}", data)

        # Get public URL
        public_url = supabase.storage.from_("templates").get_public_url(f"templates/{filename}")

        # Insert metadata into database
        supabase.table("templates").upsert({
            "filename": filename,
            "file_url": public_url,
        }).execute()

        templates = [x["filename"] for x in supabase.table("templates").select("*").execute().data]
        placeholders = ", ".join(extract_placeholders(filename)) or "No placeholders detected"

        return (
            gr.update(choices=templates, value=filename),
            gr.update(choices=templates, value=filename),
            gr.update(choices=templates, value=filename),
            placeholders,
            f"✅ Upload successful '{filename}'"
        )

    except Exception as e:
        return gr.update(), gr.update(), gr.update(), gr.update(), f"❌ Error uploading to database: {e}"

def handle_delete(name: str):
    if not name:
        return gr.update(), gr.update(), gr.update(), "⚠️ No template selected"
    try:
        # Check if file exists in storage (avoid errors)
        file_path = f"templates/{name}"
        list_resp = supabase.storage.from_("templates").list("templates")
        existing_files = [f["name"] for f in list_resp if isinstance(f, dict) and "name" in f]
        if name not in existing_files:
            print(f"⚠️ File {name} not found in database.")
        else:
            del_resp = supabase.storage.from_("templates").remove([file_path])
            print("🗑️ Storage delete:", del_resp)

        # Delete database record
        db_resp = supabase.table("templates").delete().eq("filename", name).execute()
        print("🗑️ Table delete:", db_resp)

        # Update dropdowns
        templates = list_templates()
        return (
            gr.update(choices=templates, value=None),
            gr.update(choices=templates, value=None),
            gr.update(choices=templates, value=None),
            f"✅ Delete successful '{name}'"
        )

    except Exception as e:
        print("❌ Delete error:", e)
        return (gr.update(),gr.update(),gr.update(),f"❌ Error deleting template: {e}")

def load_paste(text: str) -> str:
    global CACHED_DATA, CACHED_COLUMNS
    rows, errors = parse_pasted_text(text)
    if not rows: return "❌ No valid data. " + "; ".join(errors)
    CACHED_DATA = rows
    CACHED_COLUMNS = sorted({k for r in rows for k in r})
    return f"✅ Loaded {len(rows)} rows. Columns: {', '.join(CACHED_COLUMNS)}"

def load_file(upload) -> str:
    global CACHED_DATA, CACHED_COLUMNS
    if not upload:
        return "❌ No file uploaded"

    # If user selected a saved data file (string name)
    if isinstance(upload, str) and upload.lower().endswith((".xlsx", ".csv")):
        try:
            res = supabase.table("data").select("file_url").eq("filename", upload).execute()
            if not res.data:
                return f"❌ File '{upload}' not found in database."
            file_url = res.data[0]["file_url"]
            from io import BytesIO
            r = requests.get(file_url)
            if r.status_code != 200:
                return f"❌ Failed to download file (HTTP {r.status_code})"
            ext = os.path.splitext(upload)[1].lower()
            df = pd.read_csv(BytesIO(r.content)) if ext == ".csv" else pd.read_excel(BytesIO(r.content))
        except Exception as e:
            return f"❌ Error reading Supabase file: {e}"
    else:
        # Local upload
        path = upload.name if hasattr(upload, "name") else str(upload)
        try:
            df = pd.read_csv(path) if path.endswith(".csv") else pd.read_excel(path)
        except Exception as e:
            return f"❌ Error reading local file: {e}"

    df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]
    def format_cell(x):
        if pd.isna(x): return ""
        if isinstance(x, (datetime, pd.Timestamp)):
            return x.strftime("%#d %B %Y") if os.name == "nt" else x.strftime("%-d %B %Y")
        return str(x).strip()

    CACHED_DATA = df.applymap(format_cell).to_dict(orient="records")
    CACHED_COLUMNS = sorted({k for r in CACHED_DATA for k in r})
    return f"✅ Loaded {len(CACHED_DATA)} rows. Columns: {', '.join(CACHED_COLUMNS)}"

def gen_sample(template, pattern):
    if not template: return None, "❌ Select a template"
    if not CACHED_DATA:
        if data_tpl.value:  # if user selected Excel file from dropdown
            result = load_saved_excel(data_tpl.value)
            msg = result[-1]
            if not msg.startswith("✅"):
                return None, msg
    if not CACHED_DATA: return None, "❌ Load data first"
    row = dict(CACHED_DATA[0])
    path = generate_single_docx(template, row, pattern)
    return path, f"✅ Sample generated ({os.path.basename(path)})"

def gen_all(template, pattern):
    if not template:
        return None, "❌ Select a template"
    if not CACHED_DATA:
        if data_tpl.value:  # if user selected Excel file from dropdown
            result = load_saved_excel(data_tpl.value)
            msg = result[-1]
            if not msg.startswith("✅"):
                return None, msg
    if not CACHED_DATA:
        return None, "❌ Load data first"

    out_files = []
    tmp_dir = tempfile.mkdtemp()  # Create temp folder

    for row in CACHED_DATA:
        row = dict(row)
        out_files.append(generate_single_docx(template, row, pattern))

    # Create zip in same temp folder
    zip_path = os.path.join(tmp_dir, f"letters_{uuid.uuid4().hex[:6]}.zip")
    with zipfile.ZipFile(zip_path, "w") as z:
        for f in out_files:
            z.write(f, arcname=os.path.basename(f))

    return zip_path, f"✅ {len(out_files)} letters generated (Download below)"

# -------------------------
# Viva Letters Generator
# -------------------------
def generate_viva_letters(rename_prefix: Optional[str] = None):
    global STUDENT_DATA
    if not STUDENT_DATA:
        return None, "⚠️ No students loaded."

    tmp_dir = tempfile.mkdtemp()
    out_files = []
    errors = []

    for s in STUDENT_DATA:
        # --- Normalize keys ---
        s = {k.lower(): str(v).strip() for k, v in s.items() if v is not None}
        name = s.get("name", s.get("nama", "")).strip()
        tpl_choice = s.get("template", "").strip()
        program_val = s.get("program", "").strip()
        degree_val = s.get("degree", "").strip() or s.get("jenis_degree", "").strip()

        # --- Handle date ---
        date_raw = s.get("tarikh_viva", s.get("date", "")).strip()
        if not date_raw:
            date_val = datetime.now()
        else:
            try:
                date_val = pd.to_datetime(date_raw, errors="coerce")
                if pd.isna(date_val):
                    date_val = datetime.now()
            except Exception:
                date_val = datetime.now()
        date_val = date_val.strftime("%#d %B %Y") if os.name == "nt" else date_val.strftime("%-d %B %Y")

        if not name:
            errors.append("Missing student name.")
            continue
        if not tpl_choice:
            errors.append(f"No template selected for {name}.")
            continue

        # --- Locate template ---
        tpl_file = next((f for f in list_templates() if tpl_choice.lower() in f.lower()), None)
        if not tpl_file:
            errors.append(f"Template '{tpl_choice}' not found.")
            continue

        # --- Fill placeholders ---
        s["name"] = name
        s["nama"] = name
        s["program"] = program_val
        s["degree"] = degree_val
        s["jenis_degree"] = degree_val  # ✅ now matches Degree dropdown, not Program
        s["tarikh_submit"] = date_val
        s["tarikh"] = date_val

        tpl_path = get_template_path_from_supabase(tpl_file)
        if not tpl_path or not os.path.exists(tpl_path):
            errors.append(f"❌ Failed to download template {tpl_file} from database")
            continue
        doc = Document(tpl_path)
        fields = {k.lower(): v for k, v in s.items()}
        doc = replace_placeholders(doc, fields)

        # --- Rename pattern ---
        if rename_prefix and rename_prefix.strip():
            rename_pattern = rename_prefix
            for key, val in s.items():
                rename_pattern = rename_pattern.replace(f"{{{key}}}", val)
                rename_pattern = rename_pattern.replace(f"{{{key.upper()}}}", val.upper())
        else:
            base = os.path.splitext(tpl_file)[0]
            rename_pattern = base

        safe_name = re.sub(r"[^\w\s-]", "", rename_pattern).strip().replace(" ", "_")
        output_path = os.path.join(tmp_dir, f"{safe_name}.docx")

        try:
            doc.save(output_path)
            out_files.append(output_path)
        except Exception as e:
            errors.append(f"{name}: {e}")

    if not out_files:
        return None, f"❌ No valid letters generated.\nErrors: {'; '.join(errors)}"

    zip_path = os.path.join(tmp_dir, f"viva_letters_{uuid.uuid4().hex[:6]}.zip")
    with zipfile.ZipFile(zip_path, "w") as z:
        for f in out_files:
            z.write(f, arcname=os.path.basename(f))

    msg = f"✅ Generated {len(out_files)} viva letters."
    if errors:
        msg += f"\n⚠️ Some issues:\n" + "\n".join(errors[:5])

    return zip_path, msg

def load_excel_students(file):
    global STUDENT_DATA
    STUDENT_DATA = []

    if not file:
        return (*(gr.update(visible=False),) * 8,"⚠️ Please upload an Excel file.")

    try:
        df = pd.read_excel(file.name if hasattr(file, "name") else str(file))
        df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]

        # Accept both 'nama' or 'name'
        name_col = "nama" if "nama" in df.columns else "name" if "name" in df.columns else None
        if not name_col:
            return (
                *(gr.update(visible=False),) * 8,
                "❌ File must include a 'nama' or 'name' column."
            )

        STUDENT_DATA = []
        for _, r in df.iterrows():
            name = str(r.get(name_col, "")).strip()
            
            # --- Date handling ---
            date_val = r.get("tarikh_submit") or r.get("tarikh") or r.get("tarikh_viva") or r.get("date")
            if pd.notna(date_val):
                if isinstance(date_val, (int, float)) and not isinstance(date_val, bool):
                    try:
                        date_val = pd.Timestamp("1899-12-30") + pd.to_timedelta(float(date_val), unit="D")
                    except Exception:
                        date_val = pd.Timestamp.now()
                else:
                    parsed = pd.to_datetime(str(date_val), errors="coerce")
                    date_val = parsed if pd.notna(parsed) else pd.Timestamp.now()
            else:
                date_val = pd.Timestamp.now()
            
            # Always format as "14 October 2025"
            date_str = date_val.strftime("%#d %B %Y") if os.name == "nt" else date_val.strftime("%-d %B %Y")

            # --- Build record ---
            record = {
                str(k).lower().strip().replace(" ", "_"): str(v).strip()
                for k, v in r.items()
                if pd.notna(v)
            }

            record["name"] = record.get("name", record.get("nama", name))
            record["template"] = ""
            record["date"] = date_str
            record["tarikh_viva"] = date_str
            record["program"] = record.get("program", "")
            record["degree"] = record.get("degree", "")
            record["jenis_degree"] = record.get("degree", "")  # for placeholder {jenis_degree}

            STUDENT_DATA.append(record)

        rows = [[s["name"], s["template"], s["date"], s["program"], s["degree"]] for s in STUDENT_DATA]
        names = [s["name"] for s in STUDENT_DATA]

        return (
            gr.update(value=rows, visible=True, interactive=False),  # student_table
            gr.update(choices=names, value=None, visible=True),      # student_dropdown
            gr.update(visible=True),                                 # template_dropdown
            gr.update(visible=True),                                 # program_dropdown
            gr.update(visible=True),                                 # degree_dropdown 
            gr.update(visible=True),                                 # date_box
            gr.update(visible=True),                                 # save_btn
            gr.update(visible=True),                                 # generate_viva_btn
            gr.update(visible=True),                                 # rename_viva_box
            gr.update(visible=True),                                 # out_viva_zip
            gr.update(visible=True),                                 # data_dropdown
            
            f"✅ Loaded {len(names)} students. Now assign templates/programs, then press Save Changes."
        )


    except Exception as e:
        return (*(gr.update(visible=False),) * 10, f"❌ Error reading file: {e}")

STUDENT_DATA = []
def select_student(student_name):
    """
    Load selected student's saved values into the dropdown menus and text fields.
    Works for Viva tab: Template, Program, Degree, Date.
    """
    global STUDENT_DATA
    if not student_name:
        return (
            gr.update(value="", interactive=True, visible=True),
            gr.update(value="", interactive=True, visible=True),
            gr.update(value="", interactive=True, visible=True),
            gr.update(value="", interactive=True, visible=True)
        )

    student = next((s for s in STUDENT_DATA if str(s.get("name", s.get("nama", ""))).strip() == student_name), None)
    if not student:
        return (
            gr.update(value="", interactive=True, visible=True),
            gr.update(value="", interactive=True, visible=True),
            gr.update(value="", interactive=True, visible=True),
            gr.update(value="", interactive=True, visible=True)
        )

    template_val = str(student.get("template", "")).strip()
    program_val = str(student.get("program", "")).strip()
    degree_val = str(student.get("degree", student.get("jenis_degree", ""))).strip()
    date_val = str(student.get("tarikh_viva") or student.get("date") or "").strip()
    if not date_val or date_val.lower() == "nat":
        date_val = datetime.now().strftime("%Y-%m-%d")

    return (
        gr.update(value=template_val, interactive=True, visible=True),
        gr.update(value=program_val, interactive=True, visible=True),
        gr.update(value=degree_val, interactive=True, visible=True),
        gr.update(value=date_val, interactive=True, visible=True)
    )

def save_student(name, tpl, prog, degree, date):
    if not name:
        return gr.update(), "⚠️ Select a student first."
    # s = next((x for x in STUDENT_DATA if x["name"] == name), None)
    # new safe version
    s = next(
        (
            x
            for x in STUDENT_DATA
            if (
                x.get("name") == name
                or x.get("Name") == name
                or x.get("Student Name") == name
                or x.get("Nama") == name
                or x.get("nama") == name
            )
        ),
        None,
    )
    if s:
        s["template"] = tpl
        s["program"] = prog
        s["degree"] = degree
        s["jenis_degree"] = degree  # link degree dropdown to placeholder
        s["date"] = date
        s["tarikh_viva"] = date
    rows = [[
        x.get("name") or x.get("nama") or "",
        x.get("template", ""),
        x.get("date", ""),
        x.get("program", ""),
        x.get("degree", "")
    ] for x in STUDENT_DATA]
    return gr.update(value=rows), f"✅ Saved {name}'s info."

# -------------------------
# Data Handlers
# -------------------------
def list_saved_data():
    try:
        res = supabase.table("data").select("filename").execute()
        return sorted([r["filename"] for r in res.data])
    except Exception as e:
        print("⚠️ Could not list Excel data:", e)
        return []

def upload_data(file):
    filename = os.path.basename(file)
    folder = "data"
    file_path = f"{folder}/{filename}"
    try:
        # Delete old file from storage if exists
        try:
            supabase.storage.from_(folder).remove([filename])
        except Exception:
            pass

        # Upload file to storage
        with open(file, "rb") as f:
            supabase.storage.from_(folder).upload(filename, f)
        public_url = supabase.storage.from_(folder).get_public_url(filename)

        # Upsert into Supabase table "data"
        supabase.table("data").upsert({
            "filename": filename,
            "file_url": public_url
        }).execute()

        # Update dropdown list
        files = list_saved_data()
        return gr.update(choices=files, value=filename), f"✅ Upload successful '{filename}'"

    except Exception as e:
        return gr.update(), f"⚠️ Upload may have failed. ❌ {e}"

def delete_data(selected_file):
    if not selected_file:
        return gr.update(), gr.update(), gr.update(), gr.update(), "⚠️ Please select a file to delete."
    try:
        storage_resp = supabase.storage.from_("data").remove([selected_file])
        print("🗑️ Storage delete:", storage_resp)
        table_resp = supabase.table("data").delete().eq("filename", selected_file).execute()
        print("🗑️ Table delete:", table_resp)
        updated_files = list_saved_data()
        return (
            gr.update(choices=updated_files, value=None),
            gr.update(choices=updated_files, value=None),
            gr.update(choices=updated_files, value=None),
            gr.update(choices=updated_files, value=None),
            f"✅ Delete successful '{selected_file}'"
        )
    except Exception as e:
        print("❌ Delete error:", e)
        return gr.update(), gr.update(), gr.update(), gr.update(), f"❌ Error deleting {selected_file}: {str(e)}"

def preview_excel(selected_file):
    if not selected_file:
        return gr.update(visible=False), gr.update(value="❌ No file selected")

    try:
        res = supabase.table("data").select("file_url").eq("filename", selected_file).execute()
        if not res.data:
            return gr.update(visible=False), gr.update(value=f"❌ File '{selected_file}' not found in database")

        import requests
        from io import BytesIO
        r = requests.get(res.data[0]["file_url"])
        df = pd.read_excel(BytesIO(r.content))

        return gr.update(value=df.head(10), visible=True), gr.update(value="")
    except Exception as e:
        return gr.update(visible=False), gr.update(value=f"❌ Error previewing file: {e}")

def load_saved_excel(selected_file):
    global STUDENT_DATA, CACHED_DATA
    STUDENT_DATA = []
    CACHED_DATA = []

    if not selected_file:
        return (*(gr.update(visible=False),) * 10, "❌ No file selected")

    try:
        # Fetch file URL from Supabase table
        res = supabase.table("data").select("file_url").eq("filename", selected_file).execute()
        if not res.data:
            return (*(gr.update(visible=False),) * 10, f"❌ File '{selected_file}' not found in database")

        file_url = res.data[0]["file_url"]
        import requests
        from io import BytesIO

        r = requests.get(file_url)
        if r.status_code != 200:
            return (*(gr.update(visible=False),) * 10, f"❌ Failed to download file (HTTP {r.status_code})")

        # Load Excel file
        df = pd.read_excel(BytesIO(r.content))
        df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]

        # Basic validation
        if "name" not in df.columns and "nama" not in df.columns:
            return (*(gr.update(visible=False),) * 10, "❌ Excel file must include a 'Name' or 'Nama' column")

        # Convert to dict
        STUDENT_DATA = df.to_dict(orient="records")
        CACHED_DATA = STUDENT_DATA
        rows = df.values.tolist()
        student_names = [r.get("name") or r.get("nama") for r in STUDENT_DATA]
        
        # Build empty template assignment table 
        table_rows = [[ 
            r.get("name") or r.get("nama") or "", 
            "", # template 
            "", # date (user fills later) 
            r.get("program", ""), 
            r.get("degree", "") 
        ] for r in STUDENT_DATA]
        
        result = (
            gr.update(value=table_rows, visible=True, interactive=False),  # show template table
            gr.update(choices=student_names, visible=True),  # student dropdown
            gr.update(visible=True),  # template dropdown
            gr.update(visible=True),  # program dropdown
            gr.update(visible=True),  # degree dropdown
            gr.update(value=datetime.now().strftime("%Y-%m-%d"), visible=True),  # viva date
            gr.update(visible=True),  # save button
            gr.update(visible=True),  # generate button
            gr.update(visible=True),  # rename textbox
            gr.update(visible=True),  # zip output
            f"✅ Loaded '{selected_file}' with {len(df)} records"
        )
        if not isinstance(result[-1], str):
            result = (*result[:-1], str(result[-1]))
        return result

    except Exception as e:
        return (*(gr.update(visible=False),) * 10, f"❌ Error loading file: {e}")
        
def refresh_data(file):
    dropdown_update, status = upload_data(file)

    # Safely extract choices
    choices = getattr(dropdown_update, "choices", None)
    if not choices and isinstance(dropdown_update, dict):
        choices = dropdown_update.get("choices")

    if choices:
        latest_file = choices[-1]
        preview_update, preview_status = preview_excel(latest_file)
        if isinstance(preview_status, dict):
            preview_text = preview_status.get("value", "")
        else:
            preview_text = str(preview_status)
        combined_status = f"{status}\n{preview_text}".strip()
        return (dropdown_update, dropdown_update, dropdown_update, dropdown_update, preview_update, combined_status)
    else:
        return gr.update(), gr.update(), gr.update(), gr.update(), gr.update(visible=False), f"⚠️ Upload may have failed. {status}"

# -------------------------
# Chatbot
# -------------------------
# -------------------------
# Chatbot - FULLY WORKING & FAST (NO IMPORT ERRORS)
# -------------------------

# Simple rule-based + FAQ assistant – works instantly, no GPU needed
FAQ = {
    "template": "1. Open Microsoft Word\n2. Write your letter normally\n3. Use placeholders like {name}, {student_id}, {date}, {program}\n   Example: Dear {name}, your viva is on {tarikh_viva}\n4. Save as .docx → go to 'Manage Templates' → Upload",
    
    "placeholder": "Use curly braces in Word:\n• {name} → student name\n• {student_id}\n• {date} or {tarikh_viva}\n• {program}\n• {degree} or {jenis_degree}\n\nYou can also write {{name}} if you prefer double braces.",
    
    "how to rename": "In the 'Rename Files' box, type a pattern using placeholders:\n• Offer_Letter_{name}\n• Viva_{program}_{name}\n• {student_id}_Result\nLeave blank → uses template name + random code",
    
    "paste data": "Use this format (one person per line):\nname: Ahmad, student_id: A123, address: Kuala Lumpur\nname: Siti, student_id: A124, program: LT750",
    
    "viva": "Steps for viva letters:\n1. Go to 'Generate Viva Result Letters'\n2. Choose or upload Excel file\n3. Select each student → assign Template + Program + Degree\n4. Click 'Save Changes'\n5. (Optional) Type rename pattern\n6. Click 'Generate Viva Letters'",
    
    "excel": "Your Excel file needs at least one column: 'name' or 'nama'\nOptional columns:\n• program\n• degree\n• tarikh_viva (or date)\nDates will be auto-formatted as '14 October 2025'",
    
    "error": "Common errors & fixes:\n• 'Only .docx allowed' → upload Word files only\n• 'must include name/nama column' → add a column named 'name' or 'nama'\n• Template not found → upload it first in 'Manage Templates'",
    
    "date": "Just put any date in Excel (e.g. 20/10/2025, 2025-10-20, or Excel serial)\nIt will become: 20 October 2025 automatically",
    
    "sample": "Click 'Generate Sample' to see one letter before generating all.\nGreat for checking formatting!",

    "creator": "This system was created by Deliena Tasha Binti Abdul Rahim\nxdeliena on GitHub"
}

def chat_helper(message, history):
    msg = message.strip().lower()

    # Find the best matching FAQ key
    response = None
    for key, answer in FAQ.items():
        if key in msg:
            response = answer
            break

    if not response:
        response = ("I'm still learning! Here are things I can help with:\n"
                    "• how to create template\n"
                    "• placeholder format\n"
                    "• how to rename files\n"
                    "• viva letter steps\n"
                    "• excel format\n"
                    "• common errors\n"
                    "• creator\n"
                    "Just type any of those keywords!")

    # Add to history
    history.append((message, response))
    return history, ""

# -------------------------
# UI
# -------------------------
CSS = """
.small-btn button {padding:6px 10px; font-size:13px;}
textarea::placeholder {font-style: italic;}
#col_size {min-height: 5px;}

/* Keep preview boxes fixed */
#fixed_out {
    min-height: 90px;   /* always at least this tall */
    max-height: 90px;   /* prevent expansion */
    overflow-y: auto;   /* scroll if content is bigger */
}

"""
#Example of use: with gr.Column(elem_id="col_size"):

with gr.Blocks(css=CSS, title="Automated Letter System") as demo:
    gr.Markdown("# 📄 Automated Letter System")
    
    with gr.Tab("Tutorial"):
        with gr.Row():
            with gr.Column(scale=4):
                gr.Markdown("""
                ### 📥 Download Full User Manual
                Click the button below to download the complete user manual (ZIP file with detailed guides, screenshots, and examples).
                """)
        with gr.Row():
            with gr.Column():
                manual_download = gr.File(
                    label="Automated Letter System - User Manual.pdf",
                    value="Automated Letter System - User Manual.pdf",  # serves the local file directly
                    interactive=False,
                    height=100
                )
            with gr.Column(scale=2):  # This takes all remaining space
                gr.Markdown("")  # Invisible filler
               
        gr.Markdown("""
            ## How it works — quick guide
    
            1. **Manage Templates** — Upload `.docx` templates (use placeholders like `{name}`, `{address}`, `{student_id}`).
            2. **Generate Letters** — Select a template, then either paste rows of data or upload a CSV/XLSX file.
            3. **Preview** — Generate a single sample letter to check formatting.
            4. **Generate All** — Create letters for every row and download as a ZIP.
            
            **Data format (paste):**
            ```
            name: Ali, student_id: 2025A001, address: Shah Alam
            name: Siti, student_id: 2025A002, address: Johor
            ```
    
            **File renaming examples**
            - Leave `Rename Files` blank → filenames use template name + unique suffix.
            - Use `Letter_{name}` → generates `Letter_Ali.docx`.
    
    
            ---
    
            ## 📝 Creating a Template
    
            1. Open **Microsoft Word** (or equivalent programs).
            2. Write your letter normally.
            3. Wherever you want dynamic data, insert **placeholders** in curly braces:
               - `{name}` → student’s name
               - `{student_id}` → student’s ID
               - `{address}` → student’s address
               - `{date}` → auto-filled with today’s date if missing
            4. Example:
    
                ```
                Dear {name},
    
                We are pleased to inform you that your student ID is {student_id}.
                Your registered address is {address}.
    
                Date: {date}
                ```
    
            5. Save the document as a **.docx file**.
            6. Go to the **Manage Templates** tab in this app and upload it.
            7. The template will now be available in the **Generate Letters** tab.
    
            ✅ That’s it — your template is ready to use!

            ---
            ## 🎓 Viva Result Letter Guide

            1. **Upload Student Data** — Go to the **"Generate Viva Result Letters"** tab and upload an Excel file (.xlsx) with columns like:
               ```
               Name | Program | Degree | Tarikh_Viva
               Ali  | LT750   | Masters | 20 October 2025
               ```
            2. **Assign Details** — After loading, choose a student from the dropdown and assign:
               - A letter template  
               - Program (e.g., LT750, LT780)  
               - Degree (e.g., Masters, PhD)  
               - Viva date (auto-filled if in Excel)

            3. **Save Changes** — Click **Save Changes** to apply selections.

            4. **Generate Letters** — When all students are ready, enter a rename pattern (optional) like:
               - `Viva_Result_{name}`
               - `Viva_{program}_{name}`
               Then click **Generate Viva Letters** to download all letters as a ZIP file.

            ✅ Tips:
            - You can use `{name}`, `{program}`, `{degree}`, and `{tarikh_viva}` inside your Word template.
            - Dates are automatically formatted (e.g., `21 October 2025`).

        """)
        
    with gr.Tab("Generate Letters"):
        with gr.Row ():
            with gr.Column():
                gen_tpl = gr.Dropdown(label="Templates", choices=list_templates(), interactive=True)
                rename = gr.Textbox(label="Rename Files", placeholder="Offer_Letter_{name}")
                data_tpl = gr.Dropdown(label="Choose Data", choices=list_saved_data())
                with gr.Group():   # Groups them into the same box
                    paste = gr.Textbox(label="Paste data", lines=9,max_lines=9,
                    placeholder="name: Ali, student_id: 2025A001, address: Shah Alam\nname: Siti, student_id: 2025A002, address: Johor")
                    paste_btn = gr.Button("Enter Data", elem_classes="small-btn")
                with gr.Group():
                    sample_out = gr.File(label="Sample File", interactive=False)
                    sample_btn = gr.Button("Generate Sample", elem_classes="small-btn")
            
            with gr.Column():
                placeholders_box_gen = gr.Textbox(label="Placeholders", interactive=False, lines=3,max_lines=3)
                status = gr.Textbox(label="Status", interactive=False, lines=3,max_lines=3)
                with gr.Group():
                    all_out = gr.File(label="All Letters (ZIP)", interactive=False)
                    all_btn = gr.Button("Generate All", elem_classes="small-btn")
        
        #data_tpl.change(load_saved_excel, [data_tpl], [status])
        data_tpl.change(load_file, [data_tpl], [status])
        paste_btn.click(load_paste, [paste], [status])
        sample_btn.click(gen_sample, [gen_tpl, rename], [sample_out, status])
        gen_tpl.change(lambda t: ", ".join(extract_placeholders(t)) if t else "No placeholders detected",inputs=[gen_tpl],outputs=[placeholders_box_gen])
        all_btn.click(gen_all, [gen_tpl, rename], [all_out, status])

    with gr.Tab("Generate Viva Result Letters"):
        gr.Markdown("### 🎓 Viva Exam Result Letter Generator\nUpload student data, assign templates/programs, and generate all letters at once.")
        TEMPLATE_OPTIONS = list_templates()
        PROGRAM_OPTIONS = ["", "LT750", "LT780"]
        DEGREE_OPTIONS = ["", "Diploma", "Degree", "Masters", "PhD"]
    
        # Upload 
        with gr.Row():
            with gr.Column():
                data_tpl = gr.Dropdown(label="Student Data", choices=list_saved_data())
                load_excel_btn = gr.Button("Load Student Data", elem_classes="small-btn")
            with gr.Column():
                status_box = gr.Textbox(label="Status", interactive=False, lines=3)

        # Student Table (Read-only)
        student_table = gr.Dataframe(
            headers=["Name", "Template", "Date", "Program", "Degree"],
            datatype=["str", "str", "str", "str", "str"],
            row_count=(1, "dynamic"),
            label="Students (View Only)",
            interactive=False,
            visible=False
        )
    
        # Edit Student
        student_dropdown = gr.Dropdown(label="Select Student", choices=[], interactive=True, visible=False)
        template_dropdown = gr.Dropdown(label="Template", choices=TEMPLATE_OPTIONS, interactive=True, visible=False)
        program_dropdown = gr.Dropdown(label="Program", choices=PROGRAM_OPTIONS, interactive=True, visible=False)
        degree_dropdown = gr.Dropdown(label="Degree", choices=DEGREE_OPTIONS, allow_custom_value=True, interactive=True, visible=False)
        date_box = gr.Textbox(label="Date (auto-filled if available)", visible=False)
        save_btn = gr.Button("Save Changes", elem_classes="small-btn", visible=False)
    
        # Rename Pattern Box
        rename_viva_box = gr.Textbox(label="Rename Viva Letters (Optional)", placeholder="e.g., Viva_Letter_{name}", visible=False)
    
        # Generate Buttons
        with gr.Group():
            out_viva_zip = gr.File(label="Generated Viva Letters (ZIP)", interactive=False, visible=False)
            generate_viva_btn = gr.Button("Generate Viva Letters", elem_classes="small-btn", visible=False)
    
        # Logic Wiring
        load_excel_btn.click(
            load_saved_excel,
            [data_tpl],
            [
                student_table,
                student_dropdown,
                template_dropdown,
                program_dropdown,
                degree_dropdown, 
                date_box,
                save_btn,
                generate_viva_btn,
                rename_viva_box,
                out_viva_zip,  
                status_box
            ]
        )

        student_dropdown.change(select_student,[student_dropdown],[template_dropdown, program_dropdown, degree_dropdown, date_box])
        save_btn.click(save_student,[student_dropdown, template_dropdown, program_dropdown, degree_dropdown, date_box],[student_table, status_box])
        generate_viva_btn.click(generate_viva_letters,[rename_viva_box],[out_viva_zip, status_box],show_progress=True
        ).then(lambda zip_file: gr.update(visible=True, value=zip_file),[out_viva_zip],[out_viva_zip])

    with gr.Tab("Manage Data"):
        gr.Markdown("### 📂 Data Manager\nUpload, view and delete data files.")
        
        with gr.Row():
            with gr.Column():
                data_dropdown = gr.Dropdown(label="Data Files", choices=list_saved_data(), interactive=True)
        
        with gr.Row():
            with gr.Column():
                with gr.Group():
                    data_upload = gr.File(label="Upload Excel File", type="filepath")
                    data_upload_btn = gr.Button("Upload", elem_classes="small-btn")
            with gr.Column():
                data_status = gr.Textbox(label="Status", interactive=False, lines=2)
                with gr.Group():
                    delete_dropdown = gr.Dropdown(label="Delete Data", choices=list_saved_data(), interactive=True)
                    delete_btn = gr.Button("Delete", elem_classes="small-btn")

        with gr.Row():
            preview_status = gr.Textbox(label="Preview Status", interactive=False, lines=2, visible=False)
        with gr.Row():
            data_preview = gr.Dataframe(label="Excel Preview (First 10 Rows)", visible=False, interactive=False)
    
        data_upload_btn.click(refresh_data,[data_upload],[data_dropdown, delete_dropdown, data_tpl, gen_tpl, data_preview, data_status])
        delete_btn.click(delete_data,[delete_dropdown],[data_dropdown, delete_dropdown, data_tpl, gen_tpl, data_status])
        data_dropdown.change(lambda f: (*preview_excel(f), gr.update(visible=True)),[data_dropdown],[data_preview, preview_status])
    
    with gr.Tab("Manage Templates"):
        gr.Markdown("### 📂 Template Manager\n\nUpload, view placeholders and delete data files.")
        with gr.Row():
            with gr.Column():
                manage_tpl = gr.Dropdown(label="Templates", choices=list_templates(), interactive=True)
        
        with gr.Row():
            with gr.Column():
                with gr.Group():
                    up_file = gr.File(label="Upload .docx", type="filepath")
                    up_btn = gr.Button("Upload", elem_classes="small-btn")
                with gr.Group():
                    tpl_list = gr.Dropdown(label="Delete Templates", choices=list_templates(), interactive=True)
                    del_btn = gr.Button("Delete", elem_classes="small-btn")
            
            with gr.Column():
                placeholders_box = gr.Textbox(label="Placeholders", interactive=False, lines=5,max_lines=5)
                status_box = gr.Textbox(label="Status", interactive=False, lines=3, max_lines=3)

        up_btn.click(handle_upload,[up_file],[tpl_list, manage_tpl, gen_tpl, placeholders_box_gen, status_box])
        del_btn.click(handle_delete, [tpl_list], [tpl_list, manage_tpl, gen_tpl, status_box])
        tpl_list.change(lambda t: ", ".join(extract_placeholders(t)) if t else "No placeholders detected",inputs=[tpl_list],outputs=[placeholders_box])
        manage_tpl.change(lambda t: ", ".join(extract_placeholders(t)) if t else "No placeholders detected",inputs=[manage_tpl],outputs=[placeholders_box])

    with gr.Tab("Chatbot"):
        gr.Markdown("### 🤖 Chatbot\nAsk anything about using this system (templates, Excel formats, errors).")
    
        chat_history = gr.Chatbot(height=350)
        chat_input = gr.Textbox(label="Ask something...")
    
        chat_input.submit(chat_helper, [chat_input, chat_history], [chat_history, chat_input])


if __name__ == "__main__":
    demo.launch(inbrowser=True, share=True)
